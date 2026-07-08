"""
Generate a Word document summarizing thesis progress for the supervisor meeting.

Output: additional docs/Progress_Report_2026-06-21.docx

Usage:
    .venv/bin/python scripts/generate_progress_report.py
"""

from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = Path("additional docs/Progress_Report_2026-06-21.docx")
FIG_DIR = Path("additional docs/figures")


# ---------- styling helpers ----------

def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x6C)  # dark blue


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False,
             size: int = 11, align=None) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align is not None:
        p.alignment = align


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths is not None:
        for r in range(len(table.rows)):
            for c, w in enumerate(col_widths):
                table.rows[r].cells[c].width = Inches(w)

    # spacing after table
    doc.add_paragraph()


def add_figure(doc: Document, filename: str, caption: str,
               width_inches: float = 6.0) -> None:
    """Embed a figure with a centered italic caption beneath it."""
    path = FIG_DIR / filename
    if not path.exists():
        print(f"  WARN: figure not found, skipping: {path}")
        return
    doc.add_picture(str(path), width=Inches(width_inches))
    # Center the picture (last paragraph)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()  # spacing


def add_qa(doc: Document, question: str, answer_paragraphs: list[str]) -> None:
    """One Q&A block: bold Q + indented bullet-style body."""
    q = doc.add_paragraph()
    q_run = q.add_run(f"Q: {question}")
    q_run.bold = True
    q_run.font.size = Pt(11)
    q_run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x6C)

    for ap in answer_paragraphs:
        a = doc.add_paragraph()
        a.paragraph_format.left_indent = Inches(0.25)
        a_run = a.add_run(ap)
        a_run.font.size = Pt(11)

    doc.add_paragraph()  # trailing space


# ---------- main document content ----------

def build_document(doc: Document) -> None:
    # ---------------- Title page ----------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title.add_run("Master's Thesis – Progress Report")
    t.bold = True
    t.font.size = Pt(24)
    t.font.color.rgb = RGBColor(0x1F, 0x3B, 0x6C)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("Federated Learning for Audio-Visual Sensor Fusion on Edge Devices")
    s.italic = True
    s.font.size = Pt(14)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Author: Adnan Anwar Rajput\n").font.size = Pt(11)
    info.add_run("Berliner Hochschule für Technik, Department VII\n").font.size = Pt(11)
    info.add_run(f"Report date: {date.today().isoformat()}\n").font.size = Pt(11)
    info.add_run("Status: Bootstrap phase complete · Strategy comparison next").font.size = Pt(11)

    doc.add_page_break()

    # ---------------- Section 1: Project Overview ----------------
    add_heading(doc, "1. Project Overview", 1)

    add_para(doc,
        "This thesis investigates whether a robot can recognize objects by "
        "their sound when vision degrades (fog, darkness, occlusion). Two "
        "Raspberry Pi nodes equipped with a camera and a microphone collect "
        "synchronized audio-visual data; a multimodal neural network learns "
        "to map sound to object class."
    )

    add_para(doc,
        "The research compares three machine-learning strategies for "
        "training that network on the edge devices:"
    )

    add_bullets(doc, [
        "Strategy A – Centralized: each edge node uploads raw audio plus "
        "labels; the server runs the full pipeline (STFT + training).",
        "Strategy B – Hybrid STFT: each edge node performs STFT locally and "
        "uploads mel-spectrograms plus labels; the server trains.",
        "Strategy C – Federated: each edge node trains a local model on its "
        "own data; the server only averages weights (FedAvg). Raw data "
        "never leaves the device.",
    ])

    add_para(doc, "Research questions:", bold=True)
    add_numbered(doc, [
        "What are the trade-offs between bandwidth, privacy, power, "
        "accuracy and adaptation speed across the three strategies?",
        "Can audio meaningfully supplement vision when vision degrades?",
        "How quickly does each strategy adapt to previously unseen objects?",
    ])

    add_figure(doc, "01_system_architecture.png",
               "Figure 1 — Three-tier system layout: two edge Pis communicating "
               "over MQTT/Wi-Fi with a server that runs Flower (for Strategy C) "
               "or a training engine (for Strategies A and B).",
               width_inches=5.5)

    doc.add_page_break()

    # ---------------- Section 2: The Two Models ----------------
    add_heading(doc, "2. Two Different Models – Critical Distinction", 1)

    add_para(doc,
        "The project uses two neural networks. Confusing them is the most "
        "common source of misunderstanding.", italic=True
    )

    add_heading(doc, "2.1 YOLOv8 – The Labeler (never trained)", 2)
    add_bullets(doc, [
        "Pre-trained by Ultralytics on the COCO dataset (80 classes).",
        "Input: a single camera image.",
        "Output: a list of detected objects (e.g. 'car', 'dog', 'person').",
        "Role: provides automatic labels for the audio samples – this is the "
        "weak-supervision idea. Whatever the camera sees at time T becomes "
        "the label for the audio at time T.",
        "We never train YOLOv8. We use it like a calculator.",
    ])

    add_heading(doc, "2.2 Fusion Model – The Thesis Contribution (this is trained)", 2)
    add_bullets(doc, [
        "Custom architecture built for the thesis.",
        "Three branches: (1) MobileNetV3-Small for the image (frozen, "
        "pre-trained on ImageNet), (2) a custom audio CNN that processes "
        "mel-spectrograms, (3) a fusion head that combines both into a "
        "class prediction.",
        "Input: a 224×224×3 image PLUS a 51×128 mel-spectrogram.",
        "Output: predicted object class.",
        "Role: this is what the three thesis strategies actually train. "
        "It is the model whose accuracy / bandwidth / power we will compare.",
    ])

    add_heading(doc, "2.3 Quick analogy", 2)
    add_para(doc,
        "YOLO is the teacher: it already knows what objects look like and "
        "grades the audio. The fusion model is the student: it learns to "
        "recognize the same objects by their sound (with the image as a "
        "secondary clue), guided by the teacher's labels."
    )

    add_figure(doc, "02_two_models.png",
               "Figure 2 — Two distinct neural networks. YOLOv8 (left, gray) is "
               "frozen and only used to label audio samples. The fusion model "
               "(right, orange) is what the thesis actually trains.")

    add_heading(doc, "2.4 The three strategies — at a glance", 2)
    add_para(doc,
        "All three strategies use the same model architecture (the fusion "
        "model) and the same goal (classify the audio). They differ in "
        "WHERE the computation happens and WHAT moves over the network."
    )

    add_figure(doc, "03_three_strategies.png",
               "Figure 3 — Side-by-side view of the three strategies. As we move "
               "from A → C, less raw information leaves the device (better "
               "privacy and lower bandwidth) but more computation happens on "
               "the constrained edge hardware.")

    doc.add_page_break()

    # ---------------- Section 3: What's Been Done ----------------
    add_heading(doc, "3. What Has Been Done So Far", 1)

    add_para(doc,
        "The bootstrap phase (4 steps) is complete. Each step produced a "
        "re-runnable Python script and durable evidence on disk.", italic=True
    )

    # Step 1
    add_heading(doc, "Step 1 — YOLO Labeler Verified", 2)
    add_bullets(doc, [
        "Downloaded the YOLOv8n model (6.5 MB, 80 COCO classes).",
        "Tested on five reference images covering different classes.",
        "All five produced the correct dominant label "
        "(bus, person, dog, bird, horse).",
        "Inference latency: ~32 ms per image on the laptop's CPU.",
        "Fixed a real bug in the project scaffold: the class-filter "
        "argument was overriding the documented 'accept-all' behaviour.",
    ])

    # Step 2
    add_heading(doc, "Step 2 — Auto-Labelling Validated on Real Videos", 2)
    add_para(doc,
        "Before trusting YOLO's labels to train the audio classifier, we "
        "measured how often they actually match reality. We pulled 150 "
        "YouTube thumbnails from the VGGSound dataset (where humans have "
        "verified what is heard in the audio): 30 clips each of 'car "
        "passing by', 'driving motorcycle', 'dog barking', 'cat meowing' "
        "and 'bird chirping'."
    )
    add_para(doc,
        "Result: when YOLO produces a label, it agrees with the human "
        "label ~80% of the time (precision). It produces no label at all "
        "in ~50% of cases (recall is ~43%). The decisive metric for weak "
        "supervision is precision, not recall — a missed clip simply does "
        "not create a training pair, whereas a wrong label corrupts training. "
        "80% precision is well inside the typical tolerance for proceeding."
    )

    add_figure(doc, "04_auto_labeling.png",
               "Figure 4 — The auto-labeling pipeline. The camera and microphone "
               "are read at the same instant T. Whatever YOLO sees in the frame "
               "becomes the label for the audio chunk at T. This works because "
               "if a car is visually present, its sound is likely the dominant "
               "audio at that moment.")

    # Step 3
    add_heading(doc, "Step 3 — Audio Processing Pipeline (STFT)", 2)
    add_bullets(doc, [
        "Reads WAV files, resamples to 16 kHz, extracts the centre 500 ms "
        "(matches what the Pi microphone will produce at 10 fps).",
        "Generates mel-spectrograms of shape (51, 128) using librosa.",
        "Verified visually on all 10 UrbanSound8K classes — each class has "
        "a distinctive pattern (sirens sweep, jackhammers pulse, sirens "
        "have harmonic curves, etc.).",
        "Fixed a 49→51 frame-count inconsistency that affected 8 files "
        "across the scaffold (the librosa default convention adds 2 edge "
        "frames vs the scaffold's docstring math).",
    ])

    add_figure(doc, "07_spectrogram_grid.png",
               "Figure 5 — Mel-spectrograms produced by the pipeline, one "
               "sample per UrbanSound8K class. Visual inspection: clearly "
               "different acoustic 'signatures' — jackhammers pulse, sirens "
               "show smooth frequency sweeps, drilling is dense and broadband, "
               "engine_idling is uniform low-frequency texture.")

    # Step 4
    add_heading(doc, "Step 4 — Audio Classifier Trained on UrbanSound8K", 2)
    add_para(doc,
        "UrbanSound8K is a standard public dataset: 8,732 audio clips of "
        "10 urban sound classes (car-horn, dog-bark, engine-idling, siren, "
        "jackhammer, drilling, gun-shot, air-conditioner, children-playing, "
        "street-music)."
    )
    add_bullets(doc, [
        "Split: folds 1-8 for training (7,079 clips), fold 9 for validation, "
        "fold 10 held out as the test set.",
        "Data augmentation: extract 4 evenly-spaced 500 ms windows per "
        "training clip (~3.7× more training samples).",
        "Class imbalance handled with sklearn-balanced loss weights.",
        "Architecture: 3 convolutional blocks + global average pooling + "
        "dense layers (111,370 parameters – small enough for Pi training).",
        "Training: 41 epochs, 32 minutes on laptop CPU, EarlyStopping fired "
        "and restored the best checkpoint.",
    ])
    add_para(doc, "Final test accuracy: 68.6 %.", bold=True)
    add_para(doc,
        "Per-class F1 ranges from 44 % (siren – the weakest class) to "
        "89 % (jackhammer). The confusion matrix shows that the dominant "
        "errors are between acoustically similar classes (engine-idling ↔ "
        "air-conditioner, siren ↔ children-playing) – not random noise."
    )

    add_figure(doc, "08_training_curves.png",
               "Figure 6 — Training and validation curves. Train loss converges "
               "smoothly. Val loss is erratic at the start (large class weights "
               "+ initial high learning rate) but settles after the learning "
               "rate is automatically reduced. EarlyStopping fired at epoch 41 "
               "and the model from epoch 31 was kept.")

    add_figure(doc, "09_confusion_matrix.png",
               "Figure 7 — Confusion matrix on the held-out test fold "
               "(normalised by true class). Strong diagonal for jackhammer, "
               "gun_shot, drilling. The siren row is the weakest — about a "
               "third of sirens are misclassified as dog_bark, which is the "
               "main residual error to address later.",
               width_inches=5.0)

    # Step 5
    add_heading(doc, "Step 5 — Fusion Model Architecture Verified", 2)
    add_bullets(doc, [
        "Combines: MobileNetV3-Small (image branch, 940K frozen parameters) "
        "+ custom audio CNN (audio branch) + fusion head.",
        "Total parameters: 1.33 M (390K trainable).",
        "Forward pass on a paired (image, spectrogram) batch verified — the "
        "architecture builds correctly and produces 10-class softmax output.",
        "Not yet trained – that requires paired audio-visual data, which "
        "will come from either VGGSound or live Pi capture.",
    ])

    add_figure(doc, "05_fusion_architecture.png",
               "Figure 8 — Fusion model architecture. Two parallel feature "
               "extractors (vision branch frozen and pre-trained; audio branch "
               "trained from scratch) feed a small classification head that "
               "produces the final class prediction.",
               width_inches=5.5)

    doc.add_page_break()

    # ---------------- Section 4: Key Numbers ----------------
    add_heading(doc, "4. Key Numbers", 1)

    add_table(
        doc,
        ["Metric", "Value", "Note"],
        [
            ["YOLO inference latency", "~32 ms / image", "Laptop CPU (Mac M-series)"],
            ["YOLO label precision", "~80 %", "VGGSound subset, 150 clips"],
            ["YOLO label recall", "~43 %", "Missed clips → simply skipped"],
            ["AudioCNN parameters", "111,370", "Small enough for Pi training"],
            ["FusionModel parameters", "1,329,658", "940K frozen + 390K trainable"],
            ["Spectrogram shape", "(51, 128, 1)", "500 ms @ 16 kHz"],
            ["AudioCNN test accuracy", "68.6 %", "UrbanSound8K, fold 10 holdout"],
            ["Best per-class F1", "89 % (jackhammer)", "On test fold"],
            ["Worst per-class F1", "44 % (siren)", "On test fold"],
            ["Training time", "32 min", "Full pipeline, laptop CPU"],
        ],
        col_widths=[2.0, 1.5, 2.5],
    )

    doc.add_page_break()

    # ---------------- Section 5: Known Limitations ----------------
    add_heading(doc, "5. Known Limitations (Honestly Surfaced)", 1)

    add_para(doc,
        "These are real and will be flagged in the thesis writeup. None of "
        "them blocks the strategy comparison work that comes next.", italic=True
    )

    add_numbered(doc, [
        "Audio classifier accuracy is 68.6 %, below the 70-75 % target I set "
        "after the first training pass. Acceptable as a baseline because the "
        "thesis is about the RELATIVE performance of the three strategies, "
        "not about setting a record on UrbanSound8K. Published baselines "
        "with bigger models and longer audio windows reach 73-79 %.",

        "The 'siren' class is the weakest (44 % F1). Sirens have long "
        "frequency sweeps; the 500 ms window may be too short to capture a "
        "full sweep cycle. Possible mitigations: longer audio windows, more "
        "siren training data, or a dedicated siren detector.",

        "Single-fold evaluation (fold 10 only) is used now. Published "
        "papers use 10-fold cross-validation. A full 10-fold run will be "
        "executed before the thesis report (~5 hours, runnable overnight).",

        "No Raspberry Pi hardware yet. All work so far is on the laptop. The "
        "thesis-defining 'power consumption' measurements require the INA219 "
        "current sensor on a real Pi. Strategy comparison can still produce "
        "accuracy / bandwidth / latency numbers without Pi.",

        "Auto-labels are noisy (~20 %). Will be mitigated in the live "
        "pipeline by (a) a confidence threshold of 0.3-0.5, (b) temporal "
        "voting across 3 consecutive frames, (c) skipping multi-object "
        "scenes where the dominant audio source is ambiguous.",
    ])

    doc.add_page_break()

    # ---------------- Section 6: What's Next ----------------
    add_heading(doc, "6. What Comes Next", 1)

    add_heading(doc, "6.1 Immediate next phase – Strategy Implementation (4-6 weeks)", 2)
    add_para(doc,
        "Build the three strategies (A/B/C) on mock data first, using "
        "UrbanSound8K as a replay source for what the Pi will eventually "
        "send. This produces the first apples-to-apples comparison.")
    add_bullets(doc, [
        "Step 5a: in-process implementations of A/B/C (no MQTT or Flower "
        "yet) – validates the harness produces correct numbers.",
        "Step 5b: switch to real MQTT broker (mosquitto on localhost) for "
        "A/B and real Flower server for C. Now bandwidth and latency are "
        "measured against a real network stack.",
        "Step 5c: produce the comparison table — accuracy curve over "
        "training rounds, bandwidth used, training time.",
    ])

    add_heading(doc, "6.2 Pi Sensor Integration (2-3 weeks, after hardware)", 2)
    add_bullets(doc, [
        "Camera capture via picamera2.",
        "I2S MEMS microphone via ALSA.",
        "Audio-visual synchronisation using monotonic timestamps.",
        "Re-run the three strategies on real sensor data.",
    ])

    add_heading(doc, "6.3 Thesis Experiments (4 weeks)", 2)
    add_bullets(doc, [
        "Model freshness: pre-train on 8 of 10 classes; introduce the "
        "remaining 2 mid-experiment; measure time-to-first-correct "
        "prediction across strategies. This directly tests the central "
        "hypothesis that federated learning adapts faster.",
        "Non-IID: distribute classes unevenly across the two edge nodes; "
        "show how Strategy C (FedAvg) handles client drift.",
        "Power consumption: measure with the INA219 current sensor "
        "(idle, inference, training).",
        "Bandwidth: real bytes transmitted per training round.",
    ])

    add_heading(doc, "6.4 Thesis Writing (4 weeks)", 2)
    add_para(doc,
        "Results analysis, figures, tables, and the final thesis document.")

    add_para(doc, "Total remaining work: ~4-6 months.", bold=True)

    add_figure(doc, "06_timeline.png",
               "Figure 9 — Project timeline. Dark green is complete; light green "
               "is the immediate next phase that does not require hardware. "
               "Orange/amber phases are scheduled but contingent on hardware "
               "arrival.",
               width_inches=6.0)

    doc.add_page_break()

    # ---------------- Section 7: Anticipated Q&A ----------------
    add_heading(doc, "7. Anticipated Professor Questions", 1)

    add_para(doc,
        "These are the most likely questions a critical reviewer would ask, "
        "with the answer I plan to give. The answers acknowledge limits "
        "honestly rather than overselling.", italic=True
    )

    # METHODOLOGY
    add_heading(doc, "7.1 Methodology", 2)

    add_qa(doc,
        "Why is your test accuracy (68.6 %) below the published UrbanSound8K "
        "baselines of 73-79 %?",
        [
            "Three honest reasons. First, published baselines use 10-fold "
            "cross-validation and report the MEAN; I report a single fold, "
            "where natural variance is ±3-5 %. Second, papers typically use "
            "1-4 second audio windows; I use 500 ms to match what the Pi "
            "will produce in the live pipeline (one camera frame at 10 fps "
            "gets a 500 ms audio context). Third, my model has only 111K "
            "parameters – an order of magnitude smaller than published "
            "models – because it has to train on the Pi later.",

            "Importantly, the thesis is not trying to set an UrbanSound8K "
            "record. The contribution is the comparison of three strategies; "
            "absolute accuracy matters far less than the relative deltas "
            "between strategies.",
        ]
    )

    add_qa(doc,
        "If YOLO's auto-labels are 20 % wrong, won't that poison the "
        "audio classifier?",
        [
            "Weak supervision tolerates noise much better than people "
            "intuit. 80 % precision is well inside what's typical for "
            "self-supervised audio-visual training. Wrong labels are "
            "approximately random; with enough data, they average out and "
            "the model still learns the dominant signal.",

            "Three additional filters reduce noise further: (a) only "
            "accept YOLO predictions above a confidence threshold (0.3-0.5), "
            "(b) require three consecutive frames to agree on a class "
            "(temporal voting), (c) skip frames with multiple competing "
            "objects (ambiguous label source).",
        ]
    )

    add_qa(doc,
        "Why use UrbanSound8K (audio-only) for the bootstrap when the final "
        "system is audio-visual?",
        [
            "UrbanSound8K is the cleanest, smallest, best-understood "
            "audio-classification benchmark. It lets me verify that the "
            "audio branch of the fusion model actually learns. The full "
            "audio-visual fusion model will be trained later on paired "
            "data – either VGGSound (large, YouTube-based, brittle) or "
            "live Pi capture (real domain, requires hardware).",

            "This is the standard transfer-learning pattern: pre-train on "
            "clean public data, fine-tune on the deployment domain.",
        ]
    )

    add_qa(doc,
        "How exactly do you define and measure 'model freshness'?",
        [
            "Pre-train the model on 8 of the 10 UrbanSound8K classes. "
            "Start the live pipeline (or its simulation). At a known time "
            "T, introduce the remaining 2 classes into the data stream. "
            "Measure: how many seconds/minutes/hours until each strategy's "
            "model correctly classifies at least 70 % of the new-class "
            "samples on a held-out set.",

            "Expected ordering: Strategy C (federated, local training) "
            "adapts in minutes. Strategies A and B require a server "
            "training cycle, so adaptation latency is bounded by that "
            "cycle (hours).",
        ]
    )

    # RESULTS & TECHNICAL CHOICES
    add_heading(doc, "7.2 Results and Technical Choices", 2)

    add_qa(doc,
        "Why is the 'siren' class so weak (44 % F1)?",
        [
            "Two technical reasons. (1) Sirens are characterised by long "
            "frequency sweeps lasting 1-3 seconds; a 500 ms window often "
            "captures only a single tonal plateau, which the model confuses "
            "with other tonal sounds (dog-bark, children-playing). "
            "(2) UrbanSound8K's siren clips have high acoustic diversity "
            "(police, ambulance, fire-truck, civil defence) — more than the "
            "model can disentangle from limited data.",

            "Mitigations to try: longer audio windows, multi-window "
            "majority voting at inference time, or a specialised siren "
            "detector if it remains a problem for the safety-critical "
            "scenarios in the thesis.",
        ]
    )

    add_qa(doc,
        "Why MobileNetV3 for vision and a custom CNN for audio? Why not "
        "use the latest models like AST or Whisper?",
        [
            "Both Pi 4 hardware constraint and training feasibility. "
            "MobileNetV3-Small (1.5 M parameters) is designed for "
            "mobile/edge inference and has solid ImageNet pre-training. "
            "AST (~85 M parameters) and Whisper (~74 M for tiny) are "
            "too large for the Pi's 8 GB shared RAM during training.",

            "No comparable pre-trained 'lightweight' audio model exists. "
            "I considered YAMNet and PANNs as alternatives but they are "
            "either too large or not designed for fine-tuning. A custom "
            "CNN gives full control over the parameter budget.",
        ]
    )

    add_qa(doc,
        "Can a Raspberry Pi really train a neural network in Strategy C?",
        [
            "Yes, by design. The fusion model has ~390K trainable "
            "parameters (the MobileNetV3 vision backbone is frozen). "
            "Training uses batch size 8 with gradient accumulation. "
            "Local training rounds are short (1-3 epochs on ~500 samples) "
            "and happen during idle periods, not in the inference hot path.",

            "Empirical benchmarks from federated-learning literature show "
            "Pi 4 (8 GB) can train models of this scale at ~1-3 batches "
            "per second. A round of 3 epochs × ~60 batches × 0.5 s = "
            "~90 seconds per round. Acceptable for the experiment cadence.",
        ]
    )

    add_qa(doc,
        "Does federated learning provide actual privacy guarantees?",
        [
            "Honest answer: not formal differential-privacy guarantees. "
            "Federated learning prevents raw data from leaving the device, "
            "which is a meaningful privacy improvement, but the shared "
            "model weights can still leak information about training "
            "samples (membership inference, gradient inversion).",

            "The thesis measures RELATIVE privacy ordering "
            "(C > B > A is clear from what is and isn't transmitted), "
            "not formal DP. Adding noise to weights for differential "
            "privacy is a clean follow-up that could be future work, "
            "but is outside the current scope.",
        ]
    )

    # IMPLEMENTATION & CONTRIBUTION
    add_heading(doc, "7.3 Implementation and Contribution", 2)

    add_qa(doc,
        "What is the actual research contribution? This sounds applied.",
        [
            "The contribution is a quantitative trade-off analysis of "
            "three federated-learning strategies on real heterogeneous "
            "edge hardware for a multimodal (audio + vision) task. "
            "Three things make it novel:",

            "(1) Most federated-learning papers use simulated edge nodes "
            "running on a GPU server. We use actual Pi devices with "
            "constrained compute and an INA219 power sensor. "
            "(2) Multimodal sensor fusion on edge FL is rare; most edge "
            "FL works are vision-only or audio-only. "
            "(3) The application — auditory perception when vision "
            "fails — is itself an interesting robotics problem.",
        ]
    )

    add_qa(doc,
        "What if you cannot get the Raspberry Pi hardware working in time?",
        [
            "All experiments except absolute power-consumption measurement "
            "can be done in simulation. Bandwidth, latency, accuracy and "
            "model-freshness comparisons all work without the Pi (we "
            "mock the sensor input from UrbanSound8K). The Pi adds the "
            "INA219 power numbers, which would otherwise be missing.",

            "In the worst case, I report power as 'estimated from compute "
            "operations × CPU TDP' rather than direct measurement. The "
            "thesis still stands.",
        ]
    )

    add_qa(doc,
        "Why 500 ms audio chunks? Why not 1 s or longer?",
        [
            "Two constraints. (1) The audio chunk has to be paired with a "
            "camera frame. The camera runs at 10 fps, so a frame is captured "
            "every 100 ms. A 500 ms audio chunk centred on the frame gives "
            "250 ms of audio context before and after. (2) Real-time edge "
            "inference: at 500 ms, the model can run while the next chunk "
            "is being recorded.",

            "Longer windows (1-4 s) would likely improve accuracy but break "
            "the synchronisation contract and increase latency. The thesis "
            "could include an ablation on window size, but the live system "
            "is anchored to 500 ms.",
        ]
    )

    add_qa(doc,
        "How will you compare strategies fairly when their "
        "communication patterns are so different?",
        [
            "Same input data, same model architecture, same number of "
            "samples observed, same evaluation set, same wall-clock "
            "budget. The strategies differ only in WHERE STFT runs and "
            "WHAT is transmitted. Each strategy is evaluated on the same "
            "test fold after the same number of training samples have "
            "been seen.",

            "Comparison metrics are reported as a table: final accuracy, "
            "total bytes transmitted, training time, power consumed, time "
            "to learn the new class. Each strategy gets ten runs with "
            "different seeds; we report mean ± std.",
        ]
    )

    # PROCESS
    add_heading(doc, "7.4 Process and Timeline", 2)

    add_qa(doc,
        "What's the timeline to completion?",
        [
            "Bootstrap phase done (~30 % of the thesis). Remaining: "
            "strategy implementation 4-6 weeks, Pi sensor integration "
            "2-3 weeks (in parallel with strategy work once hardware "
            "arrives), experiments 4 weeks, writing 4 weeks. "
            "Total: ~4-6 months.",

            "Buffer: experiments and writing can overlap. The biggest "
            "risk is Pi hardware delays — mitigated by being able to "
            "run all comparisons in simulation if needed.",
        ]
    )

    add_qa(doc,
        "What's the biggest risk you currently see, and what's the plan?",
        [
            "The 'model freshness' experiment is the central thesis demo "
            "and the riskiest. It requires (a) a model that can learn new "
            "classes incrementally without catastrophic forgetting, "
            "(b) realistic on-line data streams, (c) reproducible timing "
            "measurements.",

            "Mitigation: I will validate the freshness experiment first "
            "on the in-process strategy harness (no MQTT, no hardware) "
            "before adding network complexity. If the basic effect is "
            "visible there, hardening it for real network and Pi is "
            "incremental work.",
        ]
    )

    doc.add_page_break()

    # ---------------- Section 8: Quick Glossary ----------------
    add_heading(doc, "8. Quick Glossary", 1)

    add_table(
        doc,
        ["Term", "Meaning"],
        [
            ["STFT", "Short-Time Fourier Transform – converts a waveform "
                     "into a 2-D time-frequency representation"],
            ["Mel-spectrogram", "STFT projected onto the mel scale, which "
                                 "matches how human hearing perceives pitch"],
            ["FedAvg", "Federated Averaging – the standard aggregation "
                       "algorithm; the server averages client model weights"],
            ["MQTT", "Lightweight publish-subscribe protocol commonly used "
                     "in IoT; how Strategies A and B move data"],
            ["Flower (flwr)", "Open-source framework for federated learning; "
                              "used for Strategy C"],
            ["YOLOv8", "Real-time object detection model; used as the "
                       "automatic labeller (never trained in this project)"],
            ["TFLite", "TensorFlow Lite – compact model format for edge "
                       "inference on devices like the Pi"],
            ["Weak supervision", "Training a model with labels that are "
                                 "noisy or generated automatically rather "
                                 "than hand-annotated"],
            ["Non-IID", "Non-independent and non-identically-distributed; "
                        "data is unevenly distributed across edge nodes; "
                        "a known challenge in federated learning"],
            ["Modality dropout", "Randomly zeroing out one input modality "
                                 "during training so the model is robust "
                                 "to missing data at inference time"],
        ],
        col_widths=[1.8, 4.2],
    )


def main() -> int:
    doc = Document()

    # Default font tweak for the whole document
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    build_document(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    size_kb = OUT.stat().st_size / 1024
    print(f"Saved {OUT} ({size_kb:.1f} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
